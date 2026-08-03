#!/usr/bin/env python3
"""Generate PagoYa Street Team Onboarding Guide v4 as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours ──────────────────────────────────────────────────────────────────
TEAL       = RGBColor(0x00, 0xC8, 0x8E)   # PagoYa green
NAVY       = RGBColor(0x0A, 0x25, 0x40)   # dark navy
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
AMBER      = RGBColor(0xF5, 0xA6, 0x23)   # warning/note highlights
RED_ALERT  = RGBColor(0xD0, 0x02, 0x1B)
LIGHT_GREY = RGBColor(0xF4, 0xF6, 0xF8)
MID_GREY   = RGBColor(0xCC, 0xCC, 0xCC)
DARK_TEXT  = RGBColor(0x1A, 0x1A, 0x2E)

FOOTER_TEXT = "PagoYa Technologies SA de CV  |  pagoyamx.com  |  soporte@pagoyamx.com  |  August 2026"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = color or (NAVY if level == 1 else DARK_TEXT)
    return p

def add_body(doc, text, bold=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(9.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(9.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(9.5)
    return p

def add_callout(doc, text, bg="E8F9F5", left_bar_color=None, icon=""):
    """Shaded callout box using a 1-col table."""
    tbl = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    run = p.add_run((icon + "  " if icon else "") + text)
    run.font.size = Pt(9.5)
    run.italic = True
    doc.add_paragraph()  # spacer

def add_warning(doc, text):
    add_callout(doc, text, bg="FFF3CD", icon="⚠️")

def add_note(doc, text):
    add_callout(doc, text, bg="E8F9F5", icon="💡")

def add_section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run("─" * 90)
    run.font.size = Pt(6)
    run.font.color.rgb = MID_GREY

def make_table(doc, headers, rows, col_widths=None, header_bg="0A2540", alt_row=True):
    cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=cols)
    tbl.style = 'Table Grid'

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        bg = "F4F6F8" if (alt_row and ri % 2 == 0) else "FFFFFF"
        for ci, val in enumerate(row):
            set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            if isinstance(val, tuple):  # (text, bold)
                run = p.add_run(val[0])
                run.bold = val[1]
                run.font.size = Pt(9)
            else:
                run = p.add_run(str(val))
                run.font.size = Pt(9)

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacer after table
    return tbl

def add_cover_page(doc):
    # Top accent bar — simulate with a shaded table
    tbl = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "0A2540")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(18)
    p.paragraph_format.left_indent  = Inches(0.3)
    r = p.add_run("PagoYa")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = WHITE
    r2 = p.add_run("  ·  Street Team Sales Rep — Onboarding Guide")
    r2.font.size = Pt(14)
    r2.font.color.rgb = TEAL

    doc.add_paragraph()

    # Subtitle block
    tbl2 = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl2)
    cell2 = tbl2.rows[0].cells[0]
    set_cell_bg(cell2, "E8F9F5")
    p2 = cell2.paragraphs[0]
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after  = Pt(10)
    p2.paragraph_format.left_indent  = Inches(0.3)
    lines = [
        ("Puerto Vallarta & Riviera Nayarit Field Team", True, 12, NAVY),
        ("\nVersion 4.0  |  August 2026  |  INTERNAL — Do not distribute to customers", False, 9, DARK_TEXT),
        ("\nPagoYa Technologies SA de CV  |  pagoyamx.com", False, 9, DARK_TEXT),
    ]
    for text, bold, size, color in lines:
        run = p2.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color

    doc.add_paragraph()

    # What changed box
    tbl3 = doc.add_table(rows=1, cols=1)
    table_no_borders(tbl3)
    cell3 = tbl3.rows[0].cells[0]
    set_cell_bg(cell3, "FFF3CD")
    p3 = cell3.paragraphs[0]
    p3.paragraph_format.space_before = Pt(8)
    p3.paragraph_format.space_after  = Pt(8)
    p3.paragraph_format.left_indent  = Inches(0.2)
    r_hdr = p3.add_run("⚡  What's new in v4 (August 2026)\n")
    r_hdr.bold = True
    r_hdr.font.size = Pt(10)
    r_hdr.font.color.rgb = NAVY
    changes = [
        "• Fee corrected: $25 MXN per transaction (was incorrectly listed as $15 MXN in v3)",
        "• $150 MXN signup bonus added — your most powerful sales pitch",
        "• Paula AI assistant section added (every user gets her free via WhatsApp)",
        "• Territory expanded: Nayarit municipalities now open for registration",
        "• ⚠️ Water bill limitation for Nayarit reps documented",
        "• Pay structure section added (was missing from v3)",
        "• Regional Lead / Field Director role defined",
        "• Company name corrected: PagoYa Technologies SA de CV",
        "• Service catalog updated: 50+ services (was listed as 26+)",
    ]
    for line in changes:
        rc = p3.add_run(line + "\n")
        rc.font.size = Pt(9)
        rc.font.color.rgb = DARK_TEXT

    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# BUILD DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(9.5)

add_cover_page(doc)

# ── 1. WELCOME ────────────────────────────────────────────────────────────────
add_heading(doc, "1. Welcome to Team PagoYa")
add_body(doc, "Congratulations on joining the PagoYa field team. You are the face of the company in your community.")
add_body(doc, "Your mission: talk to people who pay utility bills in cash or online — electricity, water, phone — and show them there is an easier, faster way with no waiting in line.")
add_body(doc, "")
add_body(doc, "Three things to know before your first day in the field:", bold=True)
add_bullet(doc, "The platform fee is $25 MXN per transaction — flat, no hidden charges, shown before the user confirms.")
add_bullet(doc, "Every new user receives a $150 MXN signup bonus credited instantly to their wallet — this is your most powerful opening pitch.")
add_bullet(doc, "Every user also gets Paula, PagoYa's AI payment assistant on WhatsApp, free. She sends reminders, helps with payments, and answers questions automatically.")
add_note(doc, "This guide gives you everything you need to make your first sale on day one. Read it, practise the demo script, and know your territory zone.")
add_section_divider(doc)

# ── 2. WHAT IS PAGOYA ─────────────────────────────────────────────────────────
add_heading(doc, "2. What is PagoYa?")
add_body(doc, "PagoYa is a mobile payment platform that lets users pay household bills (CFE, water, Telmex, Izzi, etc.) and top up mobile credit — all from their phone. Works with or without a bank account.")
add_body(doc, "")
add_body(doc, "How it works:", bold=True)

make_table(doc,
    headers=["Step", "Option A — Has bank account / card", "Option B — Cash / no bank account"],
    rows=[
        ("1. Register",
         "Open pagoyamx.com/register. Enter name, phone number, city, colonia. Confirm via WhatsApp verification code. Done in 2 minutes.",
         "Same — open pagoyamx.com/register. Enter name, phone number, city, colonia. Confirm via WhatsApp code."),
        ("2. Load balance",
         "Pay directly with debit or credit card in the app — no OXXO trip needed.",
         "Go to any OXXO store. Show the barcode generated in the app. Pay cash. Min $50 MXN."),
        ("3. Pay the bill",
         "Select the service, enter account number, confirm. Payment processes in seconds.",
         "Same — select service, enter account number, confirm. Payment processes in seconds."),
        ("4. Receipt + Bonus",
         "Confirmation + receipt sent via WhatsApp. $150 MXN bonus credited on first payment.",
         "Same — confirmation + receipt via WhatsApp. $150 MXN bonus credited on first payment."),
    ],
    col_widths=[1.0, 2.8, 2.8],
)

add_callout(doc,
    "Key facts: Flat $25 MXN fee per transaction. No monthly fee. No subscription. No surprises. "
    "$150 MXN signup bonus credited to wallet on first payment — reps should lead with this.",
    bg="E8F9F5", icon="✅")
add_section_divider(doc)

# ── 3. SERVICES ───────────────────────────────────────────────────────────────
add_heading(doc, "3. Services PagoYa Can Pay")

make_table(doc,
    headers=["Category", "Examples"],
    rows=[
        ("Electricity", "CFE (all of Mexico)"),
        ("Landline / Internet", "Telmex, Izzi, Totalplay, Megacable, Sky, Dish, Maxcom, Starlink"),
        ("Mobile top-ups", "Telcel, AT&T, Movistar — multiple denominations ($10–$200)"),
        ("Water", "SIAPA Jalisco (Guadalajara area), SEAPAL Vallarta (Puerto Vallarta)\n⚠️ Nayarit water utilities (CAPA) NOT currently available — see note below"),
        ("Gift cards", "Netflix, Spotify, Google Play, Xbox, Amazon — fixed denominations"),
        ("50+ services total", "Gas, streaming, pay-TV, and more in the catalog"),
    ],
    col_widths=[2.0, 4.6],
)

add_warning(doc,
    "NAYARIT WATER BILLS — IMPORTANT FOR RIVIERA NAYARIT REPS: The Nayarit municipal water "
    "utility (CAPA — Comisión de Agua Potable y Alcantarillado de Nayarit) is NOT currently "
    "available in the PagoYa catalog. Do NOT promise Nayarit water bill payments. "
    "CFE (electricity), Telmex, mobile top-ups, and all national services work normally in Nayarit. "
    "Water coverage for Nayarit is on the product roadmap.")
add_section_divider(doc)

# ── 4. YOUR ROLE ──────────────────────────────────────────────────────────────
add_heading(doc, "4. Your Role as a Field Rep")
add_body(doc, "Your job is to identify prospects and guide them through the registration process. You are not a collector or technician — you are an advisor helping people save time and avoid lines.")
add_body(doc, "")
add_body(doc, "Who to talk to:", bold=True)
add_bullet(doc, "People heading to OXXO or convenience stores to pay bills in cash")
add_bullet(doc, "People who pay bills online but find it frustrating or slow")
add_bullet(doc, "Small business owners who pay business utility bills")
add_bullet(doc, "Older adults who have a phone but are not comfortable with banking apps")
add_bullet(doc, "Workers in markets, restaurants, and service industry who need quick mobile top-ups")
add_body(doc, "")
add_body(doc, "What you do NOT do:", bold=True)
add_bullet(doc, "You do not handle customer money — PagoYa never gives you cash or asks you to collect it")
add_bullet(doc, "You are not responsible for processing payments — the app does everything")
add_bullet(doc, "You do not make promises about features not in the app or timelines you cannot guarantee")
add_section_divider(doc)

# ── 5. TERRITORY ──────────────────────────────────────────────────────────────
add_heading(doc, "5. Territory Coverage")
add_body(doc, "PagoYa is now open for registration across Puerto Vallarta (Jalisco) and Riviera Nayarit. Both regions are active — reps should operate in their assigned zone only.")

add_heading(doc, "Puerto Vallarta Zones", level=2)
make_table(doc,
    headers=["Zone / Neighborhood", "Why It Works", "Best Time"],
    rows=[
        ("Zona Romántica / Col. E. Zapata",
         "Highest density of workers, renters, small business owners. Many pay CFE and Telmex monthly. High foot traffic on Basilio Badillo and Insurgentes.",
         "10am – 2pm"),
        ("5 de Diciembre",
         "Mix of local families and digital nomads. Many unbanked and semi-banked residents. Close to market areas with high cash usage.",
         "9am – 1pm"),
        ("Versalles",
         "Business district. Workers paying utilities at lunch. Service industry employees. Taquería and abarrote strip has high dwell time.",
         "12pm – 3pm"),
        ("Pitillal",
         "Densely populated local neighborhood. Very low banking penetration. High cash economy. Strongest underbanked use case in PV.",
         "9am – 12pm"),
        ("Fluvial Vallarta",
         "Large residential colonia, many families. Mix of banked and unbanked. Multiple OXXO locations nearby — easy to show the deposit flow.",
         "4pm – 7pm"),
        ("Col. Lázaro Cárdenas",
         "Working-class residential. High proportion of CFE, Telmex, and water bill payers. Low app literacy — go slowly, demo patiently.",
         "Morning"),
        ("Marina Vallarta (workers)",
         "Hotel and marina employees. Banked users who still find bill pay friction frustrating. Card payment pitch works best here.",
         "Early morning / after shift"),
        ("Mercado Río Cuale",
         "Vendors, stall owners, market workers. Cash-first economy. Strong OXXO top-up use case. Many pay Telcel recargas weekly.",
         "8am – 11am"),
    ],
    col_widths=[1.8, 3.6, 1.2],
)
add_note(doc, "Pitillal and Zona Romántica are your highest-priority PV zones. Start there in your first week.")

add_heading(doc, "Riviera Nayarit Zones (New — August 2026)", level=2)
make_table(doc,
    headers=["Zone / Municipality", "Why It Works", "Best Time", "Water Bills?"],
    rows=[
        ("Bahía de Banderas\n(Mezcales, Valle de Banderas)",
         "Large residential municipality. Mix of local workers and new residents. Strong CFE, Telmex, and mobile top-up use case.",
         "9am – 1pm",
         "⚠️ No — CAPA not available"),
        ("Bucerías",
         "High foot traffic town centre. Mix of expats and local workers. Many pay CFE and internet bills. OXXO nearby for cash loading.",
         "9am – 12pm",
         "⚠️ No — CAPA not available"),
        ("La Cruz de Huanacaxtle",
         "Marina workers, local fishing community. Growing expat population. Strong card-payment pitch for banked marina employees.",
         "Morning / early afternoon",
         "⚠️ No — CAPA not available"),
        ("Sayulita",
         "Tourist-heavy town. Mix of seasonal workers and local residents. Mobile top-ups and Telcel recargas are strong here.",
         "10am – 2pm",
         "⚠️ No — CAPA not available"),
        ("Nuevo Vallarta / Flamingos",
         "Resort corridor workers. Hotel and hospitality employees. Banked users — lead with card payment convenience pitch.",
         "Early morning / after shift",
         "⚠️ No — CAPA not available"),
        ("Tepic (city centre)",
         "State capital. Larger population, higher density. CFE, Telmex, mobile top-ups all strong. Less competitive market than PV.",
         "9am – 1pm",
         "⚠️ No — CAPA not available"),
    ],
    col_widths=[1.7, 3.0, 1.3, 1.6],
)
add_warning(doc,
    "WATER BILLS IN NAYARIT: Do not pitch water bill payment to Nayarit residents. "
    "CAPA (the Nayarit water utility) is not in the catalog. Lead with CFE (electricity), "
    "Telmex, and mobile top-ups instead — these all work normally across Nayarit. "
    "Water coverage is expected on the roadmap; your coordinator will notify you when it's live.")
add_section_divider(doc)

# ── 6. OUTREACH METHODS ───────────────────────────────────────────────────────
add_heading(doc, "6. Outreach Methods")
add_body(doc, "Use all channels available. WhatsApp follow-up and a clear one-liner are your most consistent closers.")
make_table(doc,
    headers=["Method", "How to Use It", "Best For"],
    rows=[
        ("Street walk — SE RENTA zones",
         "Walk streets with high 'Se Renta' signage. Talk to anyone entering/exiting tiendas, waiting at bus stops, or paying at OXXO.",
         "Unbanked + semi-banked users"),
        ("OXXO entrance / exit",
         "Position near the door of high-traffic OXXOs during morning bill-pay rush (9–11am). Show users there is a faster way.",
         "Cash users already in the habit"),
        ("Mercados and tianguis",
         "Approach vendors and shoppers. Demo on your phone. Leave flyers with QR. Lunch hour best.",
         "Informal economy workers"),
        ("Taquería / abarrote strip",
         "Workers on break. Short attention — lead with the one-liner: '$25 pesos, sin filas, desde tu celular. Y te damos $150 de bono.'",
         "Service industry employees"),
        ("WhatsApp follow-up",
         "If they don't register on the spot, ask for their number and send: 'Aquí el link que te comenté: pagoyamx.com — sin descargar nada. Recibes $150 MXN de bono al registrarte.'",
         "Warm leads who showed interest"),
        ("QR flyer drop",
         "Leave printed flyers at abarrotes, papelerías, lavandería waiting areas, and doctor offices. QR goes to pagoyamx.com.",
         "Passive awareness in target zones"),
        ("Referral ask",
         "After every registration: 'Conoces a alguien más que pague luz o Telmex en efectivo?' One referral per new user averages 40% conversion.",
         "Expanding warm network"),
    ],
    col_widths=[1.6, 3.4, 1.6],
)
add_section_divider(doc)

# ── 7. 5-MINUTE DEMO ──────────────────────────────────────────────────────────
add_heading(doc, "7. How to Do a 5-Minute Demo")
add_body(doc, "Practise this flow until you can do it from memory before going into the field.")
add_body(doc, "")

steps = [
    ("Step 1 — Quick greeting (30 sec)",
     '"Hi, let me introduce you to PagoYa — an app to pay your electricity bill or top up your phone without waiting in line. You get $150 pesos just for signing up. Takes 2 minutes. Can I show you?"'),
    ("Step 2 — Ask one qualifying question (15 sec)",
     '"Do you have a bank account or debit card, or do you prefer to pay in cash?" — This tells you which path to show. Both paths register the same way.'),
    ("Step 3 — Show the app (60 sec)",
     "Open pagoyamx.com on your phone. Walk through the registration screen and show the payment flow relevant to them — card payment for banked users, OXXO barcode for cash users."),
    ("Step 4 — The cost and bonus (30 sec)",
     '"The fee is a flat $25 pesos per payment — that\'s it, no monthly charge. And when you make your first payment, the $150 peso bonus you got for signing up covers it six times over."'),
    ("Step 5 — Paula (20 sec)",
     '"You also get Paula — she\'s like a personal assistant on WhatsApp. She reminds you when bills are due, helps you pay them, and answers your questions. It\'s included, completely free."'),
    ("Step 6 — Close (60 sec)",
     '"Can I register you now, or would you prefer I send you the link? Registration takes 2 minutes and the $150 pesos goes straight to your account."'),
    ("Step 7 — Follow-up",
     "If they don't register on the spot, get their number and send: pagoyamx.com — remind them about the $150 MXN bonus."),
]

for title, body in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(title + "\n")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = NAVY
    r2 = p.add_run(body)
    r2.font.size = Pt(9.5)

add_note(doc, "Your main tool: pagoyamx.com — works on any phone with internet. Nothing to download.")
add_section_divider(doc)

# ── 8. REGISTRATION ───────────────────────────────────────────────────────────
add_heading(doc, "8. How to Register a User")
add_body(doc, "Registration is 100% free and the same for banked and unbanked users.")
add_body(doc, "")

steps_reg = [
    "Open pagoyamx.com/register on the customer's phone (or yours to demonstrate).",
    "The customer enters their name, mobile number, city, and colonia.",
    "They confirm with a verification code sent via WhatsApp.",
    "They tick the WhatsApp opt-in checkbox to receive confirmations and reminders from Paula.",
    "Done — they now have a PagoYa wallet with $150 MXN bonus ready for their first payment.",
]
for i, s in enumerate(steps_reg, 1):
    add_bullet(doc, s, bold_prefix=f"{i}.")

add_body(doc, "")
add_body(doc, "Next step depends on user type:", bold=True)
make_table(doc,
    headers=["Has debit/credit card", "Has cash only"],
    rows=[
        ("They can add their card in the app and pay bills immediately — no OXXO trip needed.",
         "Direct them to any OXXO. Show them the 'Deposit at OXXO' button. They can load as little as $50 MXN to start."),
    ],
    col_widths=[3.3, 3.3],
)
add_note(doc, "Tip: If they can't register right now, share pagoyamx.com via WhatsApp and follow up the next day. Mention the $150 MXN bonus expires — creates urgency.")
add_section_divider(doc)

# ── 9. OBJECTIONS ─────────────────────────────────────────────────────────────
add_heading(doc, "9. Handling Common Objections")
make_table(doc,
    headers=["Customer says...", "You say..."],
    rows=[
        ('"I already have a bank account"',
         '"Even better — you can pay directly with your card in the app. No OXXO trip at all. And you still get the $150 pesos just for signing up."'),
        ('"I don\'t have a bank account or card"',
         '"No problem — PagoYa works without one. Just go to any OXXO with cash, get a barcode from the app, and deposit. Done."'),
        ('"What if I get overcharged?"',
         '"The fee is always a flat $25 MXN and it\'s shown before you confirm. No surprises, ever."'),
        ('"I don\'t trust giving out my info"',
         '"We only ask for your name, phone number, city, and colonia for registration. No bank details required to sign up."'),
        ('"What if my payment goes wrong?"',
         '"Every payment gets a confirmation folio. Support is at soporte@pagoyamx.com and Paula on WhatsApp can help immediately."'),
        ('"I already pay at OXXO directly"',
         '"With PagoYa you do it from your phone — no line at the register. Load once, pay anytime. Plus you get $150 pesos free."'),
        ('"I don\'t know how to use apps"',
         '"It\'s a website, not an app — nothing to download. If you can send a WhatsApp, you can use PagoYa."'),
        ('"Can I pay my water bill in Nayarit?"',
         '"CFE, Telmex, and mobile top-ups all work in Nayarit. Water utilities in Nayarit are being added — your coordinator will confirm when it\'s live."'),
        ('"Can I pay my rent with this?"',
         '"Yes — there\'s a separate platform for rent called PagoSeguro. Go to pagoseguromx.com or find it linked from pagoyamx.com. Same idea, built for landlords and tenants."'),
    ],
    col_widths=[2.2, 4.4],
)
add_section_divider(doc)

# ── 10. PAY STRUCTURE (NEW) ───────────────────────────────────────────────────
add_heading(doc, "10. Pay Structure")
add_body(doc, "PagoYa pays commissions based on the activity your registered users generate — not on registrations alone. This aligns your incentive with genuine customer activation.", space_after=6)

add_heading(doc, "Field Rep", level=2)
make_table(doc,
    headers=["Event", "Your Earnings", "Hold Period", "Notes"],
    rows=[
        ("User completes registration", "$0", "—", "Registration is free for users; no per-signup rep commission"),
        ("User makes a confirmed bill payment", "$5 MXN", "7 days", "Per payment. Tracks automatically to your rep code."),
        ("User makes 5+ confirmed payments (month)", "Bonus TBD", "—", "Performance bonus tier — confirm amount with your coordinator"),
    ],
    col_widths=[2.2, 1.2, 1.0, 2.2],
)
add_body(doc, "")
add_body(doc, "How commissions are tracked:", bold=True)
add_bullet(doc, "Every user you register is linked to your unique rep code (e.g. ENG-01).")
add_bullet(doc, "When any user you recruited makes a bill payment, $5 MXN is logged automatically against your rep code.")
add_bullet(doc, "Commissions are held for 7 days (payment reversal window), then marked payable.")
add_bullet(doc, "Your coordinator reviews totals weekly. Payout method TBD — confirm with your coordinator.")
add_body(doc, "")
add_note(doc, "Example: You recruit 20 users in a month. 10 of them each make 4 payments = 40 payments × $5 = $200 MXN commission that month. Active users who pay regularly compound your earnings over time.")

add_heading(doc, "Regional Lead / Field Director", level=2)
make_table(doc,
    headers=["Role", "Base Commission", "Team Override", "Responsibilities"],
    rows=[
        ("Regional Lead / Field Director",
         "$5 MXN per payment from own recruits (same as field rep)",
         "$2 MXN per payment from each rep on your team\n(override — pending final confirmation)",
         "Zone coordination, daily rep reporting, quality control, escalation point for field issues"),
    ],
    col_widths=[1.5, 1.5, 2.0, 1.6],
)
add_warning(doc,
    "REGIONAL LEAD STRUCTURE — PENDING CONFIRMATION: The $2 MXN team override figure above is a "
    "proposed starting rate and has not yet been formally approved. Your coordinator will confirm "
    "the final override rate, team size assignment, and reporting cadence before your first week. "
    "The base $5 MXN/payment commission is active and confirmed in the system.")

add_body(doc, "Regional Lead responsibilities include:", bold=True)
add_bullet(doc, "Daily check-ins with each rep in their zone (WhatsApp group or in-person)")
add_bullet(doc, "Reviewing registration counts and quality (no duplicate or fake accounts)")
add_bullet(doc, "Reporting aggregate zone metrics to HQ weekly (registrations, payments, issues)")
add_bullet(doc, "First escalation point for customer complaints — never let reps handle complaints directly")
add_bullet(doc, "Onboarding new reps added to the region and ensuring they have materials")
add_section_divider(doc)

# ── 11. PAULA (NEW SECTION) ───────────────────────────────────────────────────
add_heading(doc, "11. Paula — Your User's AI Assistant")
add_body(doc, "Paula is PagoYa's AI-powered payment assistant, available to every user free via WhatsApp. She is not a chatbot — she proactively reaches out, reminds users when bills may be due, and helps them complete payments.")
add_body(doc, "")
add_body(doc, "What Paula does for your users:", bold=True)
add_bullet(doc, "Sends a personal welcome message after registration")
add_bullet(doc, "Guides new users through their first payment step by step")
add_bullet(doc, "Reminds users when bills are likely due based on their payment history")
add_bullet(doc, "Answers questions about balance, payments, and services via WhatsApp")
add_bullet(doc, "Escalates urgent issues to the support team automatically")
add_body(doc, "")
add_body(doc, "Why this matters for your pitch:", bold=True)
add_bullet(doc, "Paula makes PagoYa stickier than a normal payment app — users who interact with her stay active longer")
add_bullet(doc, "Use it in your close: 'You'll get a personal assistant on WhatsApp who reminds you when your CFE bill is coming up — completely free'")
add_bullet(doc, "Users must opt in to WhatsApp messages during registration — make sure they tick the consent checkbox")
add_note(doc, "Paula only activates for users who give WhatsApp consent at registration. Walk users through the registration screen and confirm they check the WhatsApp opt-in box. Without it, Paula cannot contact them.")
add_section_divider(doc)

# ── 12. DAILY TARGETS ─────────────────────────────────────────────────────────
add_heading(doc, "12. Suggested Daily Targets (Launch Phase)")
add_body(doc, "Work in high-OXXO-density, high foot-traffic zones. Morning is best for market workers and OXXO users. Afternoon for service industry workers on break.")
make_table(doc,
    headers=["Metric", "Minimum Goal", "Ideal Goal"],
    rows=[
        ("People approached", "20", "40+"),
        ("Demos completed", "8", "15+"),
        ("Registrations", "2", "5+"),
        ("WhatsApp links shared", "10", "20+"),
    ],
    col_widths=[3.0, 1.5, 1.5],
)
add_section_divider(doc)

# ── 13. KEY FACTS ─────────────────────────────────────────────────────────────
add_heading(doc, "13. Key Facts to Know by Heart")
make_table(doc,
    headers=["Fact", "Value"],
    rows=[
        ("Platform fee", "$25 MXN flat per transaction — shown before confirmation"),
        ("Signup bonus", "$150 MXN credited to wallet on first payment"),
        ("Minimum OXXO deposit", "$50 MXN"),
        ("Balance loading (card)", "In-app — instant, no OXXO needed"),
        ("Credit time (OXXO cash)", "Minutes in most cases (max 24 hrs)"),
        ("Services available", "50+ including CFE, Telmex, Izzi, Telcel, AT&T, Netflix, Google Play, Starlink, and more"),
        ("Water bills — PV", "SIAPA Jalisco + SEAPAL Vallarta ✓ available"),
        ("Water bills — Nayarit", "⚠️ NOT available (CAPA not in catalog — coming soon)"),
        ("Your commission", "$5 MXN per confirmed bill payment by your users (7-day hold)"),
        ("Paula AI assistant", "Free for every user on WhatsApp — user must opt in at registration"),
        ("User support", "soporte@pagoyamx.com — always escalate, never handle complaints yourself"),
        ("Website", "pagoyamx.com"),
        ("Rent payments", "pagoseguromx.com (also linked from pagoyamx.com homepage)"),
    ],
    col_widths=[2.4, 4.2],
)
add_section_divider(doc)

# ── 14. DO'S AND DON'TS ───────────────────────────────────────────────────────
add_heading(doc, "14. Do's and Don'ts")
make_table(doc,
    headers=["DO ✓", "DON'T ✗"],
    rows=[
        ("Be honest about costs ($25 MXN fee)", "Quote $15 MXN — that fee is incorrect and out of date"),
        ("Lead with the $150 MXN bonus — it's your strongest opener", "Promise features that don't exist in the app"),
        ("Ask if they're banked or cash-first before starting demo", "Assume everyone needs OXXO — many in PVR and Nayarit have cards"),
        ("Confirm users tick the WhatsApp opt-in checkbox during registration", "Skip the consent step — Paula cannot contact the user without it"),
        ("Show the app on your phone if customer has no internet", "Accept money from customers under any circumstance"),
        ("Share the link via WhatsApp if they don't register on the spot", "Give your personal number as official tech support"),
        ("Follow up on warm leads who didn't complete registration", "Pressure or badger anyone who is clearly not interested"),
        ("Report your daily activity to your Regional Lead", "Share internal or other users' personal data"),
        ("Operate only in your assigned territory zone", "Operate in another rep's zone without coordination"),
        ("Escalate all complaints immediately to soporte@pagoyamx.com", "Try to handle customer disputes yourself"),
        ("Tell Nayarit users that water bills are coming soon", "Promise Nayarit water bill payment — CAPA is not in the catalog yet"),
    ],
    col_widths=[3.3, 3.3],
)
add_section_divider(doc)

# ── 15. SUPPORT MATERIALS ─────────────────────────────────────────────────────
add_heading(doc, "15. Support Materials")
add_bullet(doc, "Business cards with QR code linking to pagoyamx.com")
add_bullet(doc, "Informational flyers — printed and digital (QR leads directly to registration)")
add_bullet(doc, "Access to the field team WhatsApp group (your coordinator will add you at onboarding)")
add_bullet(doc, "This onboarding guide in digital format")
add_bullet(doc, "Demo account in the app to walk through the full payment flow")
add_bullet(doc, "Zone map with assigned territory boundaries (provided by your Regional Lead)")
add_section_divider(doc)

# ── 16. CONTACT ───────────────────────────────────────────────────────────────
add_heading(doc, "16. Contact & Support")
make_table(doc,
    headers=["Contact", "Use For"],
    rows=[
        ("soporte@pagoyamx.com", "Customer complaints, user issues, escalations"),
        ("pagoyamx.com", "Registration, all user-facing transactions"),
        ("pagoseguromx.com", "Rent payment questions"),
        ("Field team WhatsApp group", "Daily rep activity reports, zone coordination, material requests"),
        ("Your Regional Lead", "Zone assignments, commission questions, performance issues"),
    ],
    col_widths=[2.4, 4.2],
)
add_warning(doc,
    "If a user makes a formal complaint, do NOT try to handle it yourself. "
    "Escalate immediately to soporte@pagoyamx.com with the user's phone number (last 4 digits only) "
    "and a description of the issue. Do not share the user's full phone number or personal details "
    "via the field team WhatsApp group.")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "attached_assets/PagoYa_Street_Team_Onboarding_EN_v4.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
